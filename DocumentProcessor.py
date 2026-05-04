import os
from pypdf import PdfReader
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

class DocumentProcessor:
    def __init__(self, chunk_size = 1000, chunk_overlap = 200):
        # the constructor
        self.splitter = RecursiveCharacterTextSplitter(chunk_size = chunk_size, chunk_overlap = chunk_overlap)

    
    def processFile(self, filePath: str):
        #logic to decide if its a pdf or docx 
        if filePath.lower().endswith('.pdf'):
            content = self._handlePDF(filePath)
        elif filePath.lower().endswith('.docx'):
            content = self._handleDOCX(filePath)
        else:
            raise ValueError("Unsupported file type")
        
        chunks = self.splitter.split_text(content)
        return chunks
        
    def _handlePDF(self, path): # underscore is to suggest that this is private
        reader = PdfReader(path)
        full_text = ""
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text += text + "\n"
        
        return full_text
    
    def _handleDOCX(self, path):
        doc = Document(path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        
        return "\n".join(full_text)